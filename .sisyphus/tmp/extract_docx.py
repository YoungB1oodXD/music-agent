from docx import Document

print("=" * 60)
print("开题报告内容")
print("=" * 60)

doc = Document('docs/2200310720李航颖-毕业设计开题报告.docx')
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print("\n" + "=" * 60)
print("任务书内容")
print("=" * 60)

doc2 = Document('docs/2200310720+李航颖+任务书-新.docx')
for para in doc2.paragraphs:
    if para.text.strip():
        print(para.text)